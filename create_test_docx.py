"""
create_test_docx.py
===================

Erzeugt eine kleine Beispieldatei ``test_input.docx`` im Projektordner, die
zum Prüfen des Leseformatierers genutzt werden kann.

Die Datei enthält:

* normalen Fließtext über mehrere Sätze,
* einen Absatz mit einem **blau-fetten Triggerwort** ("PAUSE"),
* einen Absatz mit bereits vorhandenen **roten Wörtern**,
* eine kleine Tabelle.

Aufruf::

    python create_test_docx.py
"""

from docx import Document
from docx.shared import RGBColor


SAMPLE_PARAGRAPHS = [
    (
        "Was bleibt, sind unzählige Erinnerungen, gemeinsame Erlebnisse "
        "und gemeinsam durchgestandene Schicksalsschläge. Manchmal sind es "
        "gar nicht die großen Worte oder Gesten, an die man sich erinnert, "
        "sondern die kleinen Dinge im Alltag."
    ),
    (
        "Die Selbstverständlichkeit, mit der jemand da ist. Die Wärme in "
        "einem Blick. Der Duft einer frisch zubereiteten Lasagne am "
        "Sonntagmittag."
    ),
    (
        "Wenn zwei Menschen ein ganzes Leben miteinander teilen, dann "
        "entsteht etwas, das von außen oft gar nicht greifbar ist."
    ),
]


def add_red_run(paragraph, text: str) -> None:
    """Hängt einen roten (nicht fetten) Run an einen Absatz."""
    run = paragraph.add_run(text)
    run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)


def add_blue_bold_run(paragraph, text: str) -> None:
    """Hängt einen blau-fetten Run (Triggerwort) an einen Absatz."""
    run = paragraph.add_run(text)
    run.bold = True
    run.font.color.rgb = RGBColor(0x00, 0x00, 0xFF)


def main(path: str = "test_input.docx") -> None:
    doc = Document()
    doc.add_heading("Beispiel für den Leseformatierer", level=1)

    for text in SAMPLE_PARAGRAPHS:
        doc.add_paragraph(text)

    # Absatz mit Triggerwort
    p = doc.add_paragraph("Sie atmet tief ein und lächelt kurz. ")
    add_blue_bold_run(p, "PAUSE")
    p.add_run(" Dann beginnt sie zu erzählen, leise und klar.")

    # Absatz mit bereits roten Wörtern
    p2 = doc.add_paragraph()
    p2.add_run("Manche ")
    add_red_run(p2, "Augenblicke")
    p2.add_run(" bleiben für immer ")
    add_red_run(p2, "unvergessen")
    p2.add_run(", auch wenn sie nur kurz waren.")

    # Tabelle
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Erste Zelle mit ein paar längeren Beispielwörtern."
    table.cell(0, 1).text = "Zweite Zelle: kurze Notiz."
    table.cell(1, 0).text = "Eine zweite Reihe mit weiterem Text zur Probe."
    table.cell(1, 1).text = "Auch hier steht etwas zum Formatieren."

    doc.save(path)
    print(f"Geschrieben: {path}")


if __name__ == "__main__":
    main()
