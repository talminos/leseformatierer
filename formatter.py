"""
formatter.py
============

Kernlogik des Leseformatierers.

Ziel
----
Eine bestehende Word-Datei (.docx) so umformatieren, dass das Ergebnis dem
Stil entspricht, den man bei vorbereiteten Lese-/Vortragstexten findet:

* Manche Wörter sind **fett & schwarz** (Betonungswörter, "Stützen" beim Lesen).
* Manche Wörter sind **rot** (zusätzliche Markierungen, oft Adjektive,
  Hilfswörter, ganze kleine Phrasen).
* **Blaue, fett gesetzte Triggerwörter** (Sprecher- oder Regie-Markierungen
  wie "PAUSE", "D.", Eigennamen) bleiben unverändert.
* Bereits **rote Wörter** können auf Wunsch erhalten werden.

Wichtig
-------
* Es wird keine externe KI verwendet – alles deterministische Heuristik.
* Wir erhalten Wörter, Leerzeichen und Satzzeichen 1:1.
* Tabellenzellen werden, soweit möglich, ebenfalls verarbeitet.
* Kopf- und Fußzeilen werden bewusst ignoriert (siehe README).
"""

from __future__ import annotations

import random
import re
from dataclasses import dataclass
from typing import Iterable, List, Optional, Sequence, Tuple

from docx import Document
from copy import deepcopy

from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.shared import Cm, Pt, RGBColor
from docx.text.paragraph import Paragraph
from docx.text.run import Run

# ---------------------------------------------------------------------------
# Konstanten
# ---------------------------------------------------------------------------

# Farben, die als „blau" gewertet werden (Hex ohne #).
# Dies sind typische Word-Blautöne aus den Vorlagen.
BLUE_HEX_VALUES = {"0000FF", "0070C0", "2F5496", "002060"}

# Farben, die als „rot" gewertet werden.
RED_HEX_VALUES = {"FF0000", "C00000", "FF2D2D"}

# Grün/fett wird als Kommentar-/Frage-Markierung des Kunden geschützt.
GREEN_HEX_VALUES = {"00B050", "008000", "00843D", "00A000", "70AD47", "548235"}

# Zielfarben, die der Formatierer schreibt.
RED_RGB = RGBColor(0xFF, 0x00, 0x00)        # Rot für rot markierte Wörter
BLUE_RGB = RGBColor(0x00, 0x00, 0xFF)       # Blau (nur falls Trigger neu erzeugt)
BLACK_RGB = RGBColor(0x00, 0x00, 0x00)      # Schwarz (Standard)

# Papierformat aus Kundenwunsch: 30 cm breit x 22 cm hoch, umlaufend 0,7 cm.
PAGE_WIDTH_CM = 30.0
PAGE_HEIGHT_CM = 22.0
PAGE_MARGIN_CM = 0.7

# Sternchen-Rubriken wie "Intro: ..." sollen größer gesetzt werden.
STAR_BLOCK_TITLE_FONT_SIZE_PT = 16

# Mindestlänge eines Worts, um überhaupt für rot/fett-schwarz infrage zu kommen.
MIN_WORD_LEN_LOOSE = 3   # Modus „beispielnah"
MIN_WORD_LEN_STRICT = 4  # Modus „streng"

# Anteil der geeigneten Kandidaten in einem Satz (Richtwerte).
# Nach dem Kundenbeispiel ist Blau nicht nur „Schutzfarbe“, sondern auch
# Hauptanker. Bestehende blau/fette Wörter bleiben geschützt; bei reinem Text
# erzeugen wir zusätzlich wenige blau/fette Hauptanker.
BLUE_ANCHOR_RATIO_RANGE = (0.22, 0.28) # Hauptanker dichter, aber weiterhin geschützt
RED_RATIO_RANGE = (0.36, 0.44)         # Blickführung auf geeigneten Wörtern ab 4 Buchstaben
BLACK_BOLD_RATIO_RANGE = (0.36, 0.44)  # Stützwörter kräftiger, ohne kurze Füllwörter

# Optionale Redezeilen-/Sprecheinheiten-Layoutierung.
SPEECH_LINE_SPACING = 1.15
SPEECH_SPACE_AFTER_PT = 12
SPEECH_MIN_WORDS_PER_LINE = 5
SPEECH_MAX_WORDS_PER_LINE = 13
SPEECH_MAX_LINES_PER_BLOCK = 2

# Wörter mit 1–2 Buchstaben werden i. d. R. nicht zusätzlich formatiert.
SHORT_WORD_LEN_MAX = 2

# Satzendzeichen.
SENTENCE_END_CHARS = ".!?"

# Gedankenstrich-Varianten (nach diesen Zeichen bevorzugt rot/fett-schwarz).
DASH_CHARS = {"-", "\u2013", "\u2014"}  # -, –, —

# Wörter, die im Beispiel häufig als Blickführung (rot) funktionieren.
# Die Listen sind bewusst klein und anpassbar, damit keine „KI“-Logik nötig ist.
RED_GUIDE_WORDS = {
    "ist", "sind", "war", "wird", "bleibt", "gehört", "betrifft", "braucht",
    "muss", "schaut", "versorgt", "gibt", "nimmt", "lernt", "beibringt",
    "vergessen", "verpasst", "bewundert", "interessieren",
    "nicht", "sehr", "immer", "manchmal", "natürlich", "insbesondere",
    "ebenso", "regelmäßig", "gerne", "wohl", "kaum", "eher", "umso",
    "mittags", "samstags", "während", "nach", "dem", "den", "zum", "vom",
    "für", "mit", "euch", "ihm", "ihn", "deiner", "seine", "neuesten",
}

# Wörter, die als rhythmische Stütze gut schwarz/fett wirken.
BLACK_SUPPORT_WORDS = {
    "als", "wenn", "dazu", "trotzdem", "und", "er", "das", "eine", "ein",
    "einen", "selbst", "auch", "vielen", "sein", "sehr", "euch", "ihn",
    "schule", "hause", "haus", "schwester", "beckenrand", "sportart",
}

# Kleine Wörter, die trotz kurzer Länge markiert werden dürfen.
SHORT_MARKABLE_WORDS = {"er", "du"}

# Häufige Wörter, die nicht als blaue Hauptanker taugen.
BLUE_STOPWORDS = {
    "als", "ist", "sind", "war", "wird", "sein", "seine", "seiner", "eher",
    "und", "oder", "aber", "auch", "dann", "wenn", "dazu", "damit", "dass",
    "der", "die", "das", "dem", "den", "ein", "eine", "einen", "einem",
    "mit", "für", "von", "vom", "zum", "zur", "auf", "aus", "nach", "euch",
    "ihm", "ihn", "du", "er", "sie", "es", "so", "nur", "wohl", "kaum",
}

# Wörter, die der Kunde explizit oder exemplarisch als blaue Hauptanker erwartet.
# Diese Liste darf später problemlos erweitert werden.
PREFERRED_BLUE_ANCHORS = {
    "andreas", "karl",
    "luise", "flügel", "stelle", "messnerin", "arbeit", "familie", "liebe",
    "verbundenheit", "bergen", "dolomiten", "garten", "lägerle",
    "leidenschaft", "skifahren", "kindheit", "freundschaften",
    "einzelkind", "bodenständigkeit", "anliegen", "bursche", "bäume",
    "gelassenheit", "freunde",
}

MONTH_NAMES = {
    "januar", "februar", "märz", "maerz", "april", "mai", "juni", "juli",
    "august", "september", "oktober", "november", "dezember",
}

AGE_UNIT_WORDS = {
    "jahr", "jahre", "jahren", "jährig", "jährige", "jährigen", "jähriger",
}

TIME_UNIT_WORDS = {"uhr"}

# Triggerwort-Marker: Wenn ein Run blau UND fett ist, gilt sein Text als Trigger.
# Wir „ballen" danach nicht alles direkt mit roten Wörtern.

# Reproduzierbare, aber leicht variable Auswahl pro Lauf.
_RNG = random.Random()


# ---------------------------------------------------------------------------
# Hilfsklassen
# ---------------------------------------------------------------------------


@dataclass
class Token:
    """Ein einzelnes Stück eines Absatzes.

    type:
      - "word":  Wortzeichen (Buchstaben/Ziffern, inkl. Umlaute, Apostroph).
      - "space": Whitespace.
      - "punct": Satzzeichen / sonstige Zeichen.

    Außerdem trägt jeder Token die ursprüngliche Formatierung (geerbt vom Run),
    damit wir z. B. blau-fett-Trigger oder vorhandenes Rot erhalten können.
    """

    text: str
    type: str
    bold: bool = False
    color_hex: Optional[str] = None   # z. B. "FF0000" oder None
    italic: Optional[bool] = None
    font_name: Optional[str] = None
    font_size: Optional[int] = None    # in EMU/half-points – hier nur durchgereicht
    underline: Optional[bool] = None
    locked: bool = False               # Trigger oder vorhandenes Rot, das wir behalten


# ---------------------------------------------------------------------------
# Klassifikation einzelner Runs
# ---------------------------------------------------------------------------


def _color_hex(run: Run) -> Optional[str]:
    """Liest die Run-Schriftfarbe als Hex-String (Großbuchstaben) aus, falls vorhanden."""
    try:
        rgb = run.font.color.rgb
    except Exception:
        return None
    if rgb is None:
        return None
    return str(rgb).upper()


def is_blue(color_hex: Optional[str]) -> bool:
    """Prüft, ob eine Hex-Farbe als blau gilt."""
    if not color_hex:
        return False
    return color_hex.upper() in BLUE_HEX_VALUES


def is_red(color_hex: Optional[str]) -> bool:
    """Prüft, ob eine Hex-Farbe als rot gilt."""
    if not color_hex:
        return False
    return color_hex.upper() in RED_HEX_VALUES


def is_green(color_hex: Optional[str]) -> bool:
    """Prüft, ob eine Hex-Farbe als Kommentar-Grün gilt."""
    if not color_hex:
        return False
    return color_hex.upper() in GREEN_HEX_VALUES


def is_blue_bold_trigger(run: Run) -> bool:
    """Run ist ein Triggerwort, wenn er **fett UND in einer Blautönung** gesetzt ist."""
    if not run.bold:
        return False
    return is_blue(_color_hex(run))


def is_green_bold_comment(run: Run) -> bool:
    """Run ist ein geschützter Kommentar, wenn er fett und grün ist."""
    if not run.bold:
        return False
    return is_green(_color_hex(run))


# ---------------------------------------------------------------------------
# Tokenisierung eines Absatzes
# ---------------------------------------------------------------------------


# Wir matchen Wörter inkl. deutscher Umlaute, Apostroph und Bindestrich.
_WORD_RE = re.compile(r"[A-Za-zÄÖÜäöüß0-9](?:[A-Za-zÄÖÜäöüß0-9'’]*)")


def _split_run_text(text: str) -> List[Tuple[str, str]]:
    """Zerlegt den Text eines Runs in Tokens (Wort/Whitespace/Satzzeichen).

    Liefert eine Liste von (text, type)-Paaren, in der Reihenfolge des Originals.
    """
    out: List[Tuple[str, str]] = []
    i = 0
    while i < len(text):
        ch = text[i]
        if ch.isspace():
            j = i
            while j < len(text) and text[j].isspace():
                j += 1
            out.append((text[i:j], "space"))
            i = j
            continue
        m = _WORD_RE.match(text, i)
        if m:
            out.append((m.group(0), "word"))
            i = m.end()
            continue
        # Alles andere (Satzzeichen, Gedankenstrich, Sonderzeichen) – ein Zeichen pro Token.
        out.append((ch, "punct"))
        i += 1
    return out


def paragraph_to_tokens(paragraph: Paragraph, keep_existing_red: bool) -> List[Token]:
    """Wandelt einen Absatz in eine Token-Liste um.

    Trigger (blau + fett) und – wenn `keep_existing_red` True ist – bestehende
    rote Wörter werden als ``locked`` markiert. Solche Tokens werden nie
    überschrieben.
    """
    tokens: List[Token] = []
    for run in paragraph.runs:
        text = run.text or ""
        if not text:
            continue
        col = _color_hex(run)
        is_trigger = is_blue_bold_trigger(run)
        is_green_comment = is_green_bold_comment(run)
        is_red_run = is_red(col)
        bold = bool(run.bold)
        for piece, kind in _split_run_text(text):
            tok = Token(
                text=piece,
                type=kind,
                bold=bold,
                color_hex=col,
                italic=run.italic,
                font_name=run.font.name,
                underline=run.underline,
            )
            if kind == "word":
                if is_trigger or is_green_comment:
                    tok.locked = True
                elif keep_existing_red and is_red_run:
                    tok.locked = True
                    tok.color_hex = "FF0000"
                    tok.bold = False
            elif kind == "punct":
                # Trigger-Punkte (z. B. „PAUSE.") sind selten – wir halten sie nicht
                # zwingend fest, damit Satzzeichen ihrem Satz folgen können.
                if is_trigger or is_green_comment:
                    tok.locked = True
            # Word kann ein sichtbares Wort intern in mehrere Runs zerlegen
            # (z. B. "Heut" + "e"). Ohne Zusammenführung würde nur ein Teil des
            # Wortes markiert. Direkt benachbarte Wort-Tokens werden deshalb zu
            # einem logischen Wort zusammengezogen.
            if tokens and tokens[-1].type == "word" and tok.type == "word":
                prev = tokens[-1]
                prev.text += tok.text
                if tok.locked:
                    prev.locked = True
                    prev.bold = tok.bold
                    prev.color_hex = tok.color_hex
                    prev.italic = tok.italic
                    prev.font_name = tok.font_name
                    prev.underline = tok.underline
                continue
            tokens.append(tok)
    return tokens


# ---------------------------------------------------------------------------
# Satzgrenzen
# ---------------------------------------------------------------------------


def split_into_sentences(tokens: Sequence[Token]) -> List[List[int]]:
    """Gruppiert Token-Indizes in Sätze.

    Ein Satz endet nach einem Satzendezeichen (``. ! ?``) oder am Absatzende.
    Mehrere aufeinander folgende Endezeichen (z. B. ``…``) zählen als ein Ende.
    """
    sentences: List[List[int]] = []
    current: List[int] = []
    for idx, tok in enumerate(tokens):
        current.append(idx)
        if tok.type == "punct" and tok.text in SENTENCE_END_CHARS:
            # Folge ggf. weiterer Satzendzeichen schlucken
            sentences.append(current)
            current = []
    if current:
        sentences.append(current)
    return sentences


# ---------------------------------------------------------------------------
# Kandidatenauswahl
# ---------------------------------------------------------------------------


def _eligible_word(tok: Token, min_len: int) -> bool:
    """Wort ist Kandidat, wenn es lang genug und nicht gesperrt ist."""
    if tok.type != "word" or tok.locked:
        return False
    if _is_year_or_number(tok.text):
        return True
    if tok.text[:1].isupper() and len(tok.text) >= 3 and _word_key(tok) not in BLUE_STOPWORDS:
        return True
    return len(tok.text) >= min_len


def _word_key(tok_or_text: Token | str) -> str:
    """Normalisierte Wortform für heuristische Listen."""
    text = tok_or_text.text if isinstance(tok_or_text, Token) else tok_or_text
    return text.strip().lower()


def _is_year_or_number(text: str) -> bool:
    """Jahreszahlen und Zahlen sollen als sichere blaue Trigger möglich sein."""
    cleaned = re.sub(r"[^\d]", "", text)
    if not cleaned:
        return False
    if len(cleaned) == 4 and 1800 <= int(cleaned) <= 2099:
        return True
    return len(cleaned) >= 2 and text.replace(".", "").replace(",", "").isdigit()


def _is_month_name(text: str) -> bool:
    """Erkennt ausgeschriebene Monatsnamen."""
    return _word_key(text) in MONTH_NAMES


def _is_age_unit_word(text: str) -> bool:
    """Erkennt Wörter, die zusammen mit Zahlen Altersangaben bilden."""
    return _word_key(text) in AGE_UNIT_WORDS


def _is_time_unit_word(text: str) -> bool:
    """Erkennt Uhrzeit-Einheiten."""
    return _word_key(text) in TIME_UNIT_WORDS


def _looks_like_name(tok: Token, token_idx: int, sentence_start_idx: Optional[int]) -> bool:
    """Einfache Namens-Heuristik ohne KI.

    Namen sind meist großgeschrieben. Am Satzanfang ist das unsicherer, daher
    bekommen Satzanfänge nur dann Namensgewicht, wenn sie kurz/eindeutig wirken
    oder in der bevorzugten Ankerliste stehen.
    """
    key = _word_key(tok)
    if key in BLUE_STOPWORDS:
        return False
    if not tok.text[:1].isupper():
        return False
    if tok.text.isupper() and len(tok.text) <= 3:
        return True
    if len(tok.text) < 3:
        return False
    if sentence_start_idx is not None and token_idx == sentence_start_idx:
        # Satzanfang: nur vorsichtig, sonst wird jedes erste Wort blau.
        return key in PREFERRED_BLUE_ANCHORS or 3 <= len(tok.text) <= 8
    return True


def _is_blue_anchor_candidate(tok: Token, min_len: int) -> bool:
    """Schätzt, ob ein Wort als blau/fetter Hauptanker geeignet ist.

    Hauptanker sind im Kundenbeispiel vor allem Nomen, Eigennamen,
    zusammengesetzte Begriffe und lange inhaltstragende Wörter.
    """
    if tok.type != "word" or tok.locked:
        return False
    key = _word_key(tok)
    if key in PREFERRED_BLUE_ANCHORS:
        return True
    if _is_year_or_number(tok.text):
        return True
    if tok.text[:1].isupper() and len(tok.text) >= 3 and key not in BLUE_STOPWORDS:
        return True
    if key in BLUE_STOPWORDS:
        return False
    if len(tok.text) < max(min_len, 4):
        return False
    if tok.text.isupper() and len(tok.text) >= 3:
        return True
    if tok.text[:1].isupper() and len(tok.text) >= 4:
        return True
    if len(tok.text) >= 8:
        return True
    return False


def classify_candidates(
    tokens: List[Token],
    sentence_indices: List[int],
    min_len: int,
    rng: random.Random,
) -> None:
    """Markiert in einem Satz Wörter als 'red' bzw. 'black_bold'.

    Schreibt direkt zurück: ``tok.color_hex`` und ``tok.bold`` für
    nicht-gesperrte Wort-Tokens.

    Heuristik (siehe README):

    * 25–40 % der geeigneten Wörter werden rot.
    * 20–35 % werden schwarz/fett.
    * Restliche Wörter bleiben „normal" (schwarz, nicht fett).
    * Sehr kurze Sätze (≤ 4 Kandidaten) bekommen nur 1–2 Zusatzformatierungen.
    * Satzanfang bevorzugt schwarz/fett.
    * Worte direkt nach einem Gedankenstrich bevorzugt rot oder schwarz/fett.
    * Nach einem blau-fett-Triggerwort wird die nächste Ballung gemieden.
    * Zwei rote Wörter direkt hintereinander vermeiden, wenn möglich.
    """
    all_word_idxs = [
        i for i in sentence_indices
        if tokens[i].type == "word" and not tokens[i].locked
    ]
    word_idxs = [i for i in sentence_indices if _eligible_word(tokens[i], min_len)]
    n = len(word_idxs)
    if n == 0:
        return

    blue_candidates = [i for i in word_idxs if _is_blue_anchor_candidate(tokens[i], min_len)]

    # Sehr kurze Sätze: nur 1–2 Zusatzformatierungen plus ggf. ein Hauptanker.
    if n <= 4:
        target_blue = 1 if blue_candidates and n >= 3 else 0
        target_red = 1 if n >= 2 else 0
        target_black_bold = 1 if n >= 1 else 0
    else:
        blue_ratio = rng.uniform(*BLUE_ANCHOR_RATIO_RANGE)
        red_ratio = rng.uniform(*RED_RATIO_RANGE)
        bb_ratio = rng.uniform(*BLACK_BOLD_RATIO_RANGE)
        target_blue = min(len(blue_candidates), max(1, int(round(n * blue_ratio))))
        target_red = max(1, int(round(n * red_ratio)))
        target_black_bold = max(1, int(round(n * bb_ratio)))
        if target_blue + target_red + target_black_bold > n:
            overflow = target_blue + target_red + target_black_bold - n
            target_black_bold = max(0, target_black_bold - overflow)

    # Kontext-Bonus: nach Gedankenstrich oder am Satzanfang.
    sentence_start_idx = sentence_indices[0] if sentence_indices else None
    after_dash: set[int] = set()
    after_trigger: set[int] = set()
    saw_dash = False
    saw_trigger = False
    for idx in sentence_indices:
        tok = tokens[idx]
        if tok.type == "word":
            if saw_dash:
                after_dash.add(idx)
                saw_dash = False
            if saw_trigger:
                after_trigger.add(idx)
                saw_trigger = False
        if tok.type == "punct" and tok.text in DASH_CHARS:
            saw_dash = True
        if tok.locked and tok.type == "word":
            # Trigger-Wort
            saw_trigger = True

    # Score je Kandidat (höher = lieber zuerst formatieren).
    def base_score(i: int) -> float:
        tok = tokens[i]
        key = _word_key(tok)
        s = float(len(tok.text))                 # längere Wörter zuerst
        if i == sentence_start_idx:
            s += 2.5                             # Satzanfang
        if i in after_dash:
            s += 2.0                             # Wort nach Gedankenstrich
        if i in after_trigger:
            s -= 1.5                             # gleich nach Trigger zurückhalten
        if key in RED_GUIDE_WORDS:
            s += 1.0
        if key in BLACK_SUPPORT_WORDS:
            s += 1.0
        return s + rng.random() * 0.6            # Variation

    def blue_score(i: int) -> float:
        tok = tokens[i]
        s = base_score(i)
        # Nomen/Eigennamen und lange zusammengesetzte Begriffe klar bevorzugen.
        if _is_year_or_number(tok.text):
            s += 8.0
        if _looks_like_name(tok, i, sentence_start_idx):
            s += 7.0
        if tok.text[:1].isupper():
            s += 4.0
        if tok.text.isupper():
            s += 2.0
        if len(tok.text) >= 10:
            s += 2.0
        if _word_key(tok) in RED_GUIDE_WORDS or _word_key(tok) in BLACK_SUPPORT_WORDS:
            s -= 3.0
        return s

    def red_score(i: int) -> float:
        tok = tokens[i]
        s = base_score(i)
        key = _word_key(tok)
        if key in RED_GUIDE_WORDS:
            s += 4.0
        # Kurze Funktions- und Bewegungswörter sollen öfter rot statt blau werden.
        if len(tok.text) <= 5:
            s += 1.1
        if tok.text[:1].isupper():
            s -= 1.2
        return s

    def black_score(i: int) -> float:
        tok = tokens[i]
        s = base_score(i)
        key = _word_key(tok)
        if key in BLACK_SUPPORT_WORDS:
            s += 4.0
        if i == sentence_start_idx:
            s += 3.0
        if tok.text[:1].isupper() and key not in BLACK_SUPPORT_WORDS:
            s -= 1.0
        return s

    ranked = sorted(word_idxs, key=base_score, reverse=True)

    # 1. Blau/fett: zentrale Hauptanker bei reinem Text erzeugen.
    blue_assigned: set[int] = set()
    def near_number_word(i: int, radius: int = 2) -> bool:
        """Prüft, ob in der Wortfolge nahe bei i eine Zahl/Jahreszahl steht."""
        if i not in all_word_idxs:
            return False
        pos = all_word_idxs.index(i)
        for other_pos in range(max(0, pos - radius), min(len(all_word_idxs), pos + radius + 1)):
            if other_pos == pos:
                continue
            other_text = tokens[all_word_idxs[other_pos]].text
            if _is_year_or_number(other_text) or other_text.isdigit():
                return True
        return False

    def near_time_unit_word(i: int, radius: int = 2) -> bool:
        """Prüft, ob nahe bei einer Zahl das Wort „Uhr“ steht."""
        if i not in all_word_idxs:
            return False
        pos = all_word_idxs.index(i)
        for other_pos in range(max(0, pos - radius), min(len(all_word_idxs), pos + radius + 1)):
            if other_pos == pos:
                continue
            if _is_time_unit_word(tokens[all_word_idxs[other_pos]].text):
                return True
        return False

    def is_clock_number(i: int) -> bool:
        """Erkennt Zahlenbestandteile von Uhrzeiten, auch einstellige Stunden."""
        tok = tokens[i]
        if tok.type != "word" or not tok.text.isdigit():
            return False
        value = int(tok.text)
        if near_time_unit_word(i, radius=2) and 0 <= value <= 59:
            return True
        # Uhrzeit mit Doppelpunkt, z. B. 9:30 oder 10:15, auch ohne „Uhr“.
        prev_tok = tokens[i - 1] if i - 1 >= 0 else None
        next_tok = tokens[i + 1] if i + 1 < len(tokens) else None
        if (
            ((prev_tok and prev_tok.type == "punct" and prev_tok.text == ":")
             or (next_tok and next_tok.type == "punct" and next_tok.text == ":"))
            and 0 <= value <= 59
        ):
            return True
        return False

    mandatory_blue = {
        i for i in all_word_idxs
        if (
            _is_year_or_number(tokens[i].text)
            or is_clock_number(i)
            or _word_key(tokens[i]) in PREFERRED_BLUE_ANCHORS
            or (_is_month_name(tokens[i].text) and near_number_word(i, radius=2))
            or (_is_age_unit_word(tokens[i].text) and near_number_word(i, radius=1))
            or (_is_time_unit_word(tokens[i].text) and near_number_word(i, radius=2))
        )
    }
    if mandatory_blue:
        target_blue = max(target_blue, len(mandatory_blue))
        blue_assigned.update(mandatory_blue)

    def has_adjacent_blue(i: int) -> bool:
        """Verhindert blaue Doppelanker direkt hintereinander."""
        if i in word_idxs:
            pos = word_idxs.index(i)
            for off in (-1, 1):
                neighbour_pos = pos + off
                if 0 <= neighbour_pos < len(word_idxs) and word_idxs[neighbour_pos] in blue_assigned:
                    return True
        return False

    for i in sorted(blue_candidates, key=blue_score, reverse=True):
        if len(blue_assigned) >= target_blue:
            break
        if i in blue_assigned:
            continue
        if has_adjacent_blue(i) and _word_key(tokens[i]) not in PREFERRED_BLUE_ANCHORS:
            continue
        blue_assigned.add(i)
    # Falls durch die Nachbarschaftsregel zu wenig Blau gesetzt wurde, nur noch
    # bevorzugte Anker auffüllen. So werden Namen/Schlüsselwörter nicht übersehen,
    # aber blaue Ballungen bleiben selten.
    for i in sorted(blue_candidates, key=blue_score, reverse=True):
        if len(blue_assigned) >= target_blue:
            break
        if i in blue_assigned:
            continue
        if _word_key(tokens[i]) in PREFERRED_BLUE_ANCHORS:
            blue_assigned.add(i)

    def has_adjacent_word(i: int, assigned: set[int]) -> bool:
        """Prüft direkte Nachbarschaft in der Wortfolge, nicht im Roh-Tokenstrom."""
        if i not in word_idxs:
            return False
        pos = word_idxs.index(i)
        for off in (-1, 1):
            neighbour_pos = pos + off
            if 0 <= neighbour_pos < len(word_idxs) and word_idxs[neighbour_pos] in assigned:
                return True
        return False

    def has_adjacent_strong(i: int, *groups: set[int]) -> bool:
        """Verhindert sichtbare Markierungsballungen direkt nebeneinander."""
        combined: set[int] = set()
        for group in groups:
            combined.update(group)
        return has_adjacent_word(i, combined)

    # 2. Schwarz/fett: Bevorzugt Satzanfang, Stützwörter und Gedankenstrich-Rhythmus.
    bb_assigned: set[int] = set()
    if sentence_start_idx in word_idxs and sentence_start_idx not in blue_assigned and target_black_bold > 0:
        bb_assigned.add(sentence_start_idx)
    for i in sorted(word_idxs, key=black_score, reverse=True):
        if len(bb_assigned) >= target_black_bold:
            break
        if i in bb_assigned or i in blue_assigned:
            continue
        if has_adjacent_word(i, bb_assigned):
            continue
        if i in after_dash:
            bb_assigned.add(i)
    for i in sorted(word_idxs, key=black_score, reverse=True):
        if len(bb_assigned) >= target_black_bold:
            break
        if i in bb_assigned or i in blue_assigned:
            continue
        if has_adjacent_word(i, bb_assigned):
            continue
        bb_assigned.add(i)

    # 3. Rot: aus den verbleibenden Kandidaten, ohne direkte Rot-Nachbarschaft.
    remaining = [
        i for i in sorted(word_idxs, key=red_score, reverse=True)
        if i not in bb_assigned and i not in blue_assigned
    ]
    red_assigned: set[int] = set()

    def red_neighbour(i: int) -> bool:
        # Prüft, ob das vorherige oder nächste Wort im Satz schon rot ist.
        for j in (i - 1, i + 1):
            if j in red_assigned:
                return True
        # Lockere Prüfung: in der Sequenz der Kandidaten
        if i in word_idxs:
            pos = word_idxs.index(i)
            for off in (-1, 1):
                p = pos + off
                if 0 <= p < len(word_idxs) and word_idxs[p] in red_assigned:
                    return True
        return False

    for i in remaining:
        if len(red_assigned) >= target_red:
            break
        if red_neighbour(i):
            continue
        red_assigned.add(i)
    # Falls wir wegen der Ballungsregeln zu wenig haben, fülle nur ohne direkte
    # Rot-Nachbarschaft auf. Andere Farben dürfen angrenzen, weil der Kunde
    # gerade starke optische Wechsel braucht.
    soft_target_red = target_red
    for i in remaining:
        if len(red_assigned) >= soft_target_red:
            break
        if i in red_assigned:
            continue
        if red_neighbour(i):
            continue
        red_assigned.add(i)

    # Abschluss-Sanierung: Keine direkte Schwarz/Fett- oder Rot-Nachbarschaft.
    # Farbliche Wechsel dürfen dicht stehen, weil sie dem Lidschlag-Prinzip helfen.
    for pos in range(len(word_idxs) - 1):
        left, right = word_idxs[pos], word_idxs[pos + 1]
        if left in bb_assigned and right in bb_assigned:
            bb_assigned.remove(right)
        if left in red_assigned and right in red_assigned:
            red_assigned.remove(right)

    # Anwenden auf die Tokens.
    for i in blue_assigned:
        tok = tokens[i]
        tok.bold = True
        tok.color_hex = "0000FF"
    for i in bb_assigned:
        tok = tokens[i]
        tok.bold = True
        tok.color_hex = "000000"
    for i in red_assigned:
        tok = tokens[i]
        tok.bold = False
        tok.color_hex = "FF0000"
    # Alle anderen, nicht gesperrten Wörter explizit normalisieren.
    for i in word_idxs:
        if i in blue_assigned or i in bb_assigned or i in red_assigned:
            continue
        tok = tokens[i]
        tok.bold = False
        tok.color_hex = None  # = Standard (i. d. R. schwarz/automatic)


# ---------------------------------------------------------------------------
# Absatz-Wiederaufbau
# ---------------------------------------------------------------------------


def _clear_paragraph_runs(paragraph: Paragraph) -> None:
    """Entfernt alle <w:r>-Elemente aus einem Absatz, behält aber die Absatz-Eigenschaften."""
    p_elem = paragraph._p
    for r in list(p_elem.findall(
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r"
    )):
        p_elem.remove(r)


def _apply_token_to_run(run: Run, tok: Token) -> None:
    """Schreibt Eigenschaften eines Tokens auf einen Run."""
    if tok.bold:
        run.bold = True
    else:
        run.bold = False
    if tok.color_hex:
        try:
            run.font.color.rgb = RGBColor.from_string(tok.color_hex)
        except Exception:
            pass
    # Andere Eigenschaften (italic, underline, font_name) behalten wir, wenn möglich.
    if tok.italic:
        run.italic = True
    if tok.underline:
        run.underline = True
    if tok.font_name:
        try:
            run.font.name = tok.font_name
        except Exception:
            pass


def _add_tokens_to_paragraph(paragraph: Paragraph, tokens: Sequence[Token]) -> None:
    """Hängt Tokens als formatierte Runs an einen Absatz an."""
    def signature(t: Token) -> tuple:
        # Whitespace und Satzzeichen erben Formatierung des umgebenden Worts nicht.
        if t.type != "word":
            # Whitespace neutral; Satzzeichen ggf. von Trigger-Locks dekoriert.
            if t.locked:
                return ("locked-punct", t.bold, t.color_hex)
            return ("plain",)
        return ("word", t.bold, t.color_hex, t.italic, t.underline, t.font_name)

    current_sig: Optional[tuple] = None
    current_run: Optional[Run] = None
    buffer: List[Token] = []

    def flush() -> None:
        nonlocal current_run, buffer
        if not buffer:
            return
        text = "".join(t.text for t in buffer)
        run = paragraph.add_run(text)
        # Eigenschaften vom ersten formatierten Token übernehmen
        for t in buffer:
            if t.type == "word" or t.locked:
                _apply_token_to_run(run, t)
                break
        buffer = []

    for tok in tokens:
        sig = signature(tok)
        if sig != current_sig:
            flush()
            current_sig = sig
        buffer.append(tok)
    flush()


def rebuild_paragraph(paragraph: Paragraph, tokens: Sequence[Token]) -> None:
    """Baut den Absatz aus der Token-Liste neu auf.

    Aufeinanderfolgende Tokens mit identischer Formatierung werden zu einem
    Run zusammengefasst, damit das Ergebnis kompakt bleibt.
    """
    _clear_paragraph_runs(paragraph)
    _add_tokens_to_paragraph(paragraph, tokens)


def _count_words(tokens: Sequence[Token]) -> int:
    """Zählt Wort-Tokens in einer Tokenfolge."""
    return sum(1 for tok in tokens if tok.type == "word")


def _trim_outer_spaces(tokens: Sequence[Token]) -> List[Token]:
    """Entfernt reine Rand-Leerzeichen nach künstlichen Umbrüchen."""
    out = list(tokens)
    while out and out[0].type == "space":
        out.pop(0)
    while out and out[-1].type == "space":
        out.pop()
    return out


def rebuild_paragraph_with_soft_lines(paragraph: Paragraph, lines: Sequence[Sequence[Token]]) -> None:
    """Baut einen Absatz mit weichen Zeilenumbrüchen zwischen den Zeilen."""
    _clear_paragraph_runs(paragraph)
    for line_index, line_tokens in enumerate(lines):
        if line_index > 0:
            paragraph.add_run().add_break(WD_BREAK.LINE)
        _add_tokens_to_paragraph(paragraph, line_tokens)


def _apply_speech_paragraph_format(paragraph: Paragraph) -> None:
    """Setzt die vom Kunden gewünschte Sprecheinheiten-Optik."""
    paragraph.paragraph_format.line_spacing = SPEECH_LINE_SPACING
    paragraph.paragraph_format.space_after = Pt(SPEECH_SPACE_AFTER_PT)


def _apply_manuscript_paragraph_format(paragraph: Paragraph) -> None:
    """Setzt das Grundlayout aus der Wunschdatei: 1,15 und 12 pt Abstand."""
    paragraph.paragraph_format.line_spacing = SPEECH_LINE_SPACING
    paragraph.paragraph_format.space_after = Pt(SPEECH_SPACE_AFTER_PT)


def _apply_document_page_setup(doc: Document) -> None:
    """Setzt Papierformat und Seitenränder nach Kundenwunsch."""
    for section in doc.sections:
        section.page_width = Cm(PAGE_WIDTH_CM)
        section.page_height = Cm(PAGE_HEIGHT_CM)
        section.top_margin = Cm(PAGE_MARGIN_CM)
        section.bottom_margin = Cm(PAGE_MARGIN_CM)
        section.left_margin = Cm(PAGE_MARGIN_CM)
        section.right_margin = Cm(PAGE_MARGIN_CM)


def _is_star_separator(text: str) -> bool:
    """Erkennt reine Sternchen-Trennzeilen."""
    stripped = text.strip()
    return len(stripped) >= 5 and set(stripped) == {"*"}


def _find_star_block_indices(paragraphs: Sequence[Paragraph]) -> set[int]:
    """Findet Absätze, die zu *...*-Einschüben gehören.

    Beispiel:
    *****************************************************
    Intro: Andreas Gabalier…
    *****************************************************

    Der komplette Block wird von der normalen Wortmarkierung ausgenommen.
    """
    indices: set[int] = set()
    i = 0
    while i < len(paragraphs):
        if not _is_star_separator(paragraphs[i].text):
            i += 1
            continue
        end = None
        for j in range(i + 1, min(len(paragraphs), i + 8)):
            if _is_star_separator(paragraphs[j].text):
                end = j
                break
        if end is not None:
            indices.update(range(i, end + 1))
            i = end + 1
        else:
            indices.add(i)
            i += 1
    return indices


def _apply_star_block_format(paragraph: Paragraph) -> None:
    """Formatiert Sternchen-Einschübe: einfach, fett, mittig.

    Die eigentliche Rubrikzeile zwischen den Sternchen wird zusätzlich auf
    16 pt gesetzt; die Sternchen-Trennlinien bleiben in der bestehenden Größe.
    """
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    for run in paragraph.runs:
        run.bold = True
        if not _is_star_separator(paragraph.text):
            run.font.size = Pt(STAR_BLOCK_TITLE_FONT_SIZE_PT)
        # Farbe bewusst nicht überschreiben, falls der Nutzer dort etwas markiert hat.


def _insert_paragraph_after(paragraph: Paragraph) -> Paragraph:
    """Fügt direkt nach einem Absatz einen neuen Absatz mit kopierten Absatzformaten ein."""
    new_p = deepcopy(paragraph._p)
    # Inhalt entfernen, Absatzformat aber behalten.
    for child in list(new_p):
        if child.tag.endswith("}r"):
            new_p.remove(child)
    paragraph._p.addnext(new_p)
    return Paragraph(new_p, paragraph._parent)


def _token_is_break_opportunity(tok: Token) -> bool:
    """Geeignete Trennstellen für Sprecheinheiten."""
    return tok.type == "punct" and (tok.text in SENTENCE_END_CHARS or tok.text in DASH_CHARS)


def split_into_speech_lines(tokens: Sequence[Token]) -> List[List[Token]]:
    """Teilt einen Absatz in kurze Redezeilen.

    Word-Zeilenzahl ist ohne Layout-Engine nicht exakt berechenbar. Diese
    Heuristik erzeugt deshalb kurze Zeilen und bricht nur an semantisch
    vertretbaren Stellen: nach Gedankenstrichen oder Satzenden. Kommas bleiben
    ausdrücklich innerhalb derselben Sprecheinheit, damit kein Absatz mit
    Komma, Punkt oder einem losgelösten Nebensatz beginnt.
    """
    lines: List[List[Token]] = []
    current: List[Token] = []
    words = 0
    overdue = False

    for tok in tokens:
        current.append(tok)
        if tok.type == "word":
            words += 1
            if words >= SPEECH_MAX_WORDS_PER_LINE:
                overdue = True

        # Der Gedankenstrich und Satzenden sind natürliche Schnittstellen.
        # Auch wenn eine Zeile länger wird, warten wir auf so eine Schnittstelle,
        # statt mitten im Satz oder direkt vor/nach einem Komma zu trennen.
        if words >= SPEECH_MIN_WORDS_PER_LINE and _token_is_break_opportunity(tok):
            lines.append(_trim_outer_spaces(current))
            current = []
            words = 0
            overdue = False

        # Sicherheitsventil für extrem lange Sätze ohne jeden Gedankenstrich
        # oder Punkt: lieber eine zu lange Einheit als hässliche Komma-Brüche.
        # Erst ab sehr hoher Überlänge brechen wir nach dem nächsten Wort.
        elif overdue and words >= SPEECH_MAX_WORDS_PER_LINE * 2 and tok.type == "word":
            lines.append(_trim_outer_spaces(current))
            current = []
            words = 0
            overdue = False
    if current:
        lines.append(_trim_outer_spaces(current))
    return [line for line in lines if line]


def group_speech_lines(lines: Sequence[Sequence[Token]]) -> List[List[List[Token]]]:
    """Bündelt Redezeilen zu Absätzen.

    Je Block entstehen 1–2 sichtbare Zeilen: Die Zeilen im Block werden mit
    weichem Word-Umbruch verbunden; zwischen den Blöcken kommt ein harter
    Absatz mit 12 pt Abstand. Dadurch entsteht die Optik aus dem Screenshot,
    ohne dass jede einzelne Zeile 12 pt Abstand bekommt.
    """
    blocks: List[List[List[Token]]] = []
    current: List[List[Token]] = []
    for line in lines:
        current.append(list(line))
        if len(current) >= SPEECH_MAX_LINES_PER_BLOCK:
            blocks.append(current)
            current = []
    if current:
        blocks.append(current)
    return blocks


def rebuild_paragraph_as_speech_units(paragraph: Paragraph, tokens: Sequence[Token]) -> None:
    """Baut einen Absatz als Sprecheinheiten neu auf.

    Jede Redezeilen-Gruppe wird ein harter Word-Absatz mit 12 pt Abstand danach.
    Innerhalb dieser Gruppe stehen 1–2 Zeilen per weichem Zeilenumbruch.
    """
    lines = split_into_speech_lines(tokens)
    blocks = group_speech_lines(lines)
    if len(blocks) <= 1:
        rebuild_paragraph_with_soft_lines(paragraph, blocks[0] if blocks else [tokens])
        _apply_speech_paragraph_format(paragraph)
        return

    rebuild_paragraph_with_soft_lines(paragraph, blocks[0])
    _apply_speech_paragraph_format(paragraph)
    last = paragraph
    for block in blocks[1:]:
        last = _insert_paragraph_after(last)
        rebuild_paragraph_with_soft_lines(last, block)
        _apply_speech_paragraph_format(last)


# ---------------------------------------------------------------------------
# Haupt-Pipeline
# ---------------------------------------------------------------------------


def _process_paragraph(
    paragraph: Paragraph,
    *,
    min_len: int,
    keep_existing_red: bool,
    only_trigger_paragraphs: bool,
    speech_units: bool,
    manuscript_layout: bool,
    rng: random.Random,
) -> bool:
    """Verarbeitet einen einzelnen Absatz. Gibt True zurück, falls bearbeitet."""
    if not paragraph.runs:
        return False

    # Optional: nur Absätze bearbeiten, in denen blau-fette Trigger oder
    # bereits fett gesetzte Markierungen vorhanden sind.
    if only_trigger_paragraphs:
        has_trigger = any(is_blue_bold_trigger(r) for r in paragraph.runs)
        # „blau/fett-Trigger" – wir sind hier strikt.
        if not has_trigger:
            return False

    tokens = paragraph_to_tokens(paragraph, keep_existing_red=keep_existing_red)
    if not tokens:
        return False
    sentences = split_into_sentences(tokens)
    for sent in sentences:
        classify_candidates(tokens, sent, min_len=min_len, rng=rng)
    if speech_units:
        rebuild_paragraph_as_speech_units(paragraph, tokens)
    else:
        rebuild_paragraph(paragraph, tokens)
        if manuscript_layout:
            _apply_manuscript_paragraph_format(paragraph)
    return True


def _iter_table_paragraphs(doc: Document) -> Iterable[Paragraph]:
    """Gibt alle Absätze aus Tabellenzellen zurück (rekursiv)."""
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p
                # Verschachtelte Tabellen
                for nested in cell.tables:
                    for nrow in nested.rows:
                        for ncell in nrow.cells:
                            for p in ncell.paragraphs:
                                yield p


def format_document(
    src_path: str,
    dest_path: str,
    *,
    mode: str = "loose",
    keep_existing_red: bool = True,
    only_trigger_paragraphs: bool = False,
    speech_units: bool = False,
    manuscript_layout: bool = True,
    seed: Optional[int] = None,
) -> dict:
    """Hauptfunktion: liest ``src_path``, formatiert und speichert nach ``dest_path``.

    Parameters
    ----------
    mode :
        ``"loose"`` (beispielnah, ab 3 Buchstaben) oder ``"strict"`` (ab 4
        Buchstaben).
    keep_existing_red :
        Vorhandene rote Wörter unverändert lassen.
    only_trigger_paragraphs :
        Nur Absätze mit blau-fetten Triggerwörtern bearbeiten.
    speech_units :
        Lange Absätze zusätzlich in kurze Sprecheinheiten aufteilen und mit
        1,15 Zeilenabstand sowie 12 pt Abstand nach Absatz formatieren.
    manuscript_layout :
        Grundlayout mit 1,15 Zeilenabstand und 12 pt Abstand nach Absatz
        anwenden, auch wenn keine Redezeilen erzeugt werden.
    seed :
        Optionaler Zufallsseed für reproduzierbare Ergebnisse.
    """
    if mode == "strict":
        min_len = MIN_WORD_LEN_STRICT
    else:
        min_len = MIN_WORD_LEN_LOOSE

    rng = random.Random(seed) if seed is not None else random.Random()

    doc = Document(src_path)
    _apply_document_page_setup(doc)

    processed = 0
    star_block_indices = _find_star_block_indices(doc.paragraphs)
    # 1) Top-Level-Absätze
    for idx, paragraph in enumerate(doc.paragraphs):
        if idx in star_block_indices:
            _apply_star_block_format(paragraph)
            processed += 1
            continue
        if _process_paragraph(
            paragraph,
            min_len=min_len,
            keep_existing_red=keep_existing_red,
            only_trigger_paragraphs=only_trigger_paragraphs,
            speech_units=speech_units,
            manuscript_layout=manuscript_layout,
            rng=rng,
        ):
            processed += 1
    # 2) Tabellenzellen
    for paragraph in _iter_table_paragraphs(doc):
        if _process_paragraph(
            paragraph,
            min_len=min_len,
            keep_existing_red=keep_existing_red,
            only_trigger_paragraphs=only_trigger_paragraphs,
            speech_units=speech_units,
            manuscript_layout=manuscript_layout,
            rng=rng,
        ):
            processed += 1
    # Kopf-/Fußzeilen werden bewusst nicht angefasst (siehe README).

    doc.save(dest_path)
    return {"processed_paragraphs": processed}
